from __future__ import annotations

"""
Small Flask web app for Latmay.

Responsibilities:
1. Show the resume submission form.
2. Read form inputs.
3. Call the ranking backend.
4. Render the results page.

UI is kept in:
- templates/home.html
- templates/results.html

Ranking logic is kept in:
- ranking_service.py
"""

import os
import time
import traceback
from io import BytesIO
from pathlib import Path
from typing import Any

from flask import Flask, g, jsonify, render_template, request, send_from_directory
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.datastructures import FileStorage
from werkzeug.exceptions import HTTPException

from recent_jobs_service import get_recent_jobs_page
from recent_posted_filter import RECENT_POSTED_HOURS


def int_from_env(name: str, default: int) -> int:
    """
    Read an integer environment variable with an explicit default.
    """
    return int(os.environ.get(name, str(default)))


def bool_from_env(name: str, default: bool = False) -> bool:
    raw_value = os.environ.get(name)
    if raw_value is None:
        return default
    return raw_value.strip().lower() in {"true", "1", "yes", "on"}


app = Flask(__name__)

FIREBASE_WEB_CONFIG = {
    "apiKey": os.environ.get("FIREBASE_WEB_API_KEY", ""),
    "authDomain": "project-a674ada3-351d-46e0-8b7.firebaseapp.com",
    "projectId": "project-a674ada3-351d-46e0-8b7",
    "storageBucket": "project-a674ada3-351d-46e0-8b7.firebasestorage.app",
    "messagingSenderId": "2821343980",
    "appId": "1:2821343980:web:278fbea2cfecb90bc933c9",
}

TOP_K_TO_SHOW = int_from_env("TOP_K_TO_SHOW", 10)
MAX_REQUEST_BYTES = int_from_env("MAX_REQUEST_BYTES", 5000000)
MAX_RESUME_CHARS = int_from_env("MAX_RESUME_CHARS", 12000)
MAX_RESUME_FILE_BYTES = int_from_env("MAX_RESUME_FILE_BYTES", 4000000)
RANK_RATE_LIMIT = os.environ.get("RANK_RATE_LIMIT", "10 per minute")
SHOW_RANKING_DIAGNOSTICS = bool_from_env("SHOW_RANKING_DIAGNOSTICS", False)
ENABLE_MAX_YOE_FILTER = bool_from_env("ENABLE_MAX_YOE_FILTER", True)
SAFE_ERROR_MESSAGE_MAX_CHARS = int_from_env("SAFE_ERROR_MESSAGE_MAX_CHARS", 200)
SAFE_TRACEBACK_MAX_FRAMES = int_from_env("SAFE_TRACEBACK_MAX_FRAMES", 8)
ACCOUNT_DELETE_RECENT_AUTH_SECONDS = int_from_env("ACCOUNT_DELETE_RECENT_AUTH_SECONDS", 600)

app.config["MAX_CONTENT_LENGTH"] = MAX_REQUEST_BYTES

limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=[],
    storage_uri=os.environ.get("RATELIMIT_STORAGE_URI", "memory://"),
)

ALLOWED_RESUME_FILE_EXTENSIONS = {".pdf", ".docx"}


def verify_firebase_id_token(id_token: str) -> dict[str, Any]:
    """Verify a Firebase ID token using Admin SDK application-default credentials."""
    import firebase_admin
    from firebase_admin import auth

    if not firebase_admin._apps:
        firebase_admin.initialize_app()
    return auth.verify_id_token(id_token)


def delete_firebase_user(firebase_uid: str) -> None:
    """Permanently delete a Firebase Authentication user by verified UID."""
    import firebase_admin
    from firebase_admin import auth

    if not firebase_admin._apps:
        firebase_admin.initialize_app()
    auth.delete_user(firebase_uid)


@app.before_request
def establish_firebase_identity() -> Any:
    """Attach an optional verified Firebase identity to Flask's request context."""
    g.firebase_uid = None
    g.firebase_user = None

    authorization = request.headers.get("Authorization", "").strip()
    if not authorization:
        return None

    scheme, separator, id_token = authorization.partition(" ")
    if scheme.lower() != "bearer" or not separator or not id_token.strip():
        return jsonify({"error": "Invalid Authorization header."}), 401

    try:
        decoded_token = verify_firebase_id_token(id_token.strip())
    except Exception as exc:
        print(
            "authentication: Firebase ID token rejected: "
            f"error_type={type(exc).__name__}",
            flush=True,
        )
        return jsonify({"error": "Invalid or expired authentication token."}), 401

    uid = str(decoded_token.get("uid") or decoded_token.get("sub") or "").strip()
    if not uid:
        return jsonify({"error": "Authentication token has no user identifier."}), 401

    g.firebase_uid = uid
    g.firebase_user = decoded_token
    return None


class ResumeExtractionError(ValueError):
    """
    User-facing resume extraction failure.
    """


def safe_error_message(exc: Exception) -> str:
    """
    Return a short exception message without multi-line payloads.
    """
    message = " ".join(str(exc).split())
    if not message:
        return ""
    return message[:SAFE_ERROR_MESSAGE_MAX_CHARS]


def safe_traceback_summary(exc: Exception) -> str:
    """
    Log traceback locations only; do not include source lines or local values.
    """
    frames = traceback.extract_tb(exc.__traceback__)[-SAFE_TRACEBACK_MAX_FRAMES:]
    parts = [
        f"{os.path.basename(frame.filename)}:{frame.lineno}:{frame.name}"
        for frame in frames
    ]
    return " > ".join(parts)


def log_safe_request_error(route_name: str, exc: Exception) -> None:
    """
    Log request failures without writing resume text or exception payloads.
    """
    print(
        "ALERT: request failed: "
        f"route={route_name}, "
        f"error_type={type(exc).__name__}, "
        f"traceback={safe_traceback_summary(exc)!r}",
        flush=True,
    )


def parse_float_or_none(value: str | None) -> float | None:
    """
    Convert a form input string into a float, or None if blank.

    Used for optional numeric filters like max required years of experience.
    """
    if value is None:
        return None

    value = value.strip()

    if not value:
        return None

    return float(value)


def normalize_extracted_text(text: str) -> str:
    """
    Collapse extraction whitespace into readable plain text.
    """
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def uploaded_file_extension(uploaded_file: FileStorage) -> str:
    """
    Return the lowercase file extension for an uploaded resume.
    """
    return Path(uploaded_file.filename or "").suffix.lower()


def read_limited_upload(uploaded_file: FileStorage) -> bytes:
    """
    Read an uploaded resume and enforce the resume-file byte limit.
    """
    file_bytes = uploaded_file.read(MAX_RESUME_FILE_BYTES + 1)
    if len(file_bytes) > MAX_RESUME_FILE_BYTES:
        raise ResumeExtractionError(
            f"Resume file is too large. Please upload a file under {MAX_RESUME_FILE_BYTES / 1_000_000:g} MB."
        )
    if not file_bytes:
        raise ResumeExtractionError("Please choose a PDF or DOCX resume file.")
    return file_bytes


def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract text from a PDF resume.
    """
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(file_bytes))
    if reader.is_encrypted:
        raise ResumeExtractionError("That PDF is encrypted. Please upload an unlocked PDF or paste the resume text.")

    page_text = []
    for page in reader.pages:
        page_text.append(page.extract_text() or "")

    return normalize_extracted_text("\n".join(page_text))


def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX resume.
    """
    from docx import Document

    document = Document(BytesIO(file_bytes))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    return normalize_extracted_text("\n".join(parts))


def extract_resume_text_from_upload(uploaded_file: FileStorage) -> str:
    """
    Extract plain text from a supported resume upload.
    """
    extension = uploaded_file_extension(uploaded_file)
    if extension not in ALLOWED_RESUME_FILE_EXTENSIONS:
        raise ResumeExtractionError("Please upload a PDF or DOCX resume file.")

    file_bytes = read_limited_upload(uploaded_file)
    try:
        if extension == ".pdf":
            resume_text = extract_pdf_text(file_bytes)
        else:
            resume_text = extract_docx_text(file_bytes)
    except ResumeExtractionError:
        raise
    except Exception as exc:
        raise ResumeExtractionError("Could not extract text from that resume file. Please paste the text instead.") from exc

    if not resume_text:
        raise ResumeExtractionError("No readable text was found in that resume file. Please paste the resume text instead.")

    return resume_text


def get_rank_form_inputs(*, include_resume: bool = True) -> tuple[str, str | None, str | None, float | None, bool, bool]:
    """
    Read and clean form inputs from the /rank POST request.
    """
    resume_input_mode = request.form.get("resume_input_mode", "upload").strip().lower()
    uploaded_file = request.files.get("resume_file")
    resume_text = request.form.get("resume_text", "").strip()
    if include_resume and resume_input_mode == "upload" and uploaded_file is not None and uploaded_file.filename:
        resume_text = extract_resume_text_from_upload(uploaded_file)
    if not include_resume:
        resume_text = ""

    country = request.form.get("country", "").strip() or None
    state = request.form.get("state", "").strip() or None
    max_required_yoe = parse_float_or_none(request.form.get("max_required_yoe")) if ENABLE_MAX_YOE_FILTER else None
    exclude_security_clearance = request.form.get("exclude_security_clearance") == "on"
    require_recent_posted = request.form.get("require_recent_posted") == "on"

    return resume_text, country, state, max_required_yoe, exclude_security_clearance, require_recent_posted


def require_firebase_uid() -> str:
    uid = str(getattr(g, "firebase_uid", None) or "").strip()
    if not uid:
        raise PermissionError("Login is required for saved resume profiles.")
    return uid


def render_home(
    *,
    error: str | None = None,
    notice: str | None = None,
    resume_text: str = "",
    country: str | None = None,
    state: str | None = None,
    max_required_yoe: float | None = None,
    exclude_security_clearance: bool = False,
    require_recent_posted: bool = False,
) -> str:
    """
    Render the homepage form with current form state.
    """
    return render_template(
        "home.html",
        error=error,
        notice=notice,
        resume_text=resume_text,
        country=country,
        state=state,
        max_required_yoe=max_required_yoe,
        enable_max_yoe_filter=ENABLE_MAX_YOE_FILTER,
        exclude_security_clearance=exclude_security_clearance,
        require_recent_posted=require_recent_posted,
        recent_posted_hours=RECENT_POSTED_HOURS,
        max_resume_file_bytes=MAX_RESUME_FILE_BYTES,
        max_resume_file_mb=f"{MAX_RESUME_FILE_BYTES / 1_000_000:g}",
        firebase_web_config=FIREBASE_WEB_CONFIG,
        show_auth_controls=True,
    )


@app.errorhandler(413)
def request_too_large(_: Exception) -> tuple[str, int]:
    """
    Return a friendly page when Flask rejects an oversized request body.
    """
    return render_home(
        error=f"That request is too large. Please keep uploads under {MAX_REQUEST_BYTES:,} bytes.",
    ), 413


@app.errorhandler(429)
def rate_limit_exceeded(_: Exception) -> tuple[str, int]:
    """
    Return a friendly page when /rank is submitted too frequently.
    """
    return render_home(
        error="Too many ranking requests. Please wait a moment and try again.",
    ), 429


@app.get("/")
def home() -> str:
    """
    Show the homepage form.
    """
    return render_home()


@app.get("/api/recent-jobs")
@limiter.limit("60 per minute")
def recent_jobs() -> Any:
    """
    Return one shard's worth of recently posted jobs for the homepage feed.

    Shards are most-recent-first, so shard=0 is the newest jobs; the client
    fetches additional shards as the visitor scrolls further.
    """
    try:
        shard_index = int(request.args.get("shard", "0").strip())
    except ValueError:
        return jsonify({"error": "shard must be an integer."}), 400
    if shard_index < 0:
        return jsonify({"error": "shard must not be negative."}), 400

    try:
        page = get_recent_jobs_page(shard_index)
    except Exception as exc:
        log_safe_request_error("/api/recent-jobs", exc)
        return jsonify({"error": "Could not load recent jobs."}), 500

    if page is None:
        return jsonify({"shard": shard_index, "jobs": [], "has_more": False})

    return jsonify(page)


@app.post("/resume-preview")
def resume_preview() -> str | tuple[str, int]:
    """
    Extract text from an uploaded resume and show it for review before ranking.
    """
    try:
        uploaded_file = request.files.get("resume_file")
        if uploaded_file is None or not uploaded_file.filename:
            return render_home(error="Please choose a PDF or DOCX resume file."), 400

        resume_text = extract_resume_text_from_upload(uploaded_file)
        if len(resume_text) > MAX_RESUME_CHARS:
            return render_home(
                error=f"Extracted resume text is too long. Please edit it under {MAX_RESUME_CHARS:,} characters.",
                resume_text=resume_text[:MAX_RESUME_CHARS],
            ), 413

        return render_home(
            notice="Resume text extracted. Review it below, make any edits, then rank jobs.",
            resume_text=resume_text,
        )
    except ResumeExtractionError as exc:
        return render_home(error=str(exc)), 400
    except HTTPException:
        raise
    except Exception as exc:
        log_safe_request_error("/resume-preview", exc)
        return render_home(error="Something went wrong while extracting that resume. Please paste the text instead."), 500


@app.post("/rank")
@limiter.limit(RANK_RATE_LIMIT)
def rank() -> str | tuple[str, int]:
    """
    Rank jobs using the submitted resume and filters.
    """
    try:
        use_saved_resume = request.form.get("use_saved_resume") == "on"
        (
            resume_text,
            country,
            state,
            max_required_yoe,
            exclude_security_clearance,
            require_recent_posted,
        ) = get_rank_form_inputs(include_resume=not use_saved_resume)

        cached_resume_profile = None
        if use_saved_resume:
            try:
                firebase_uid = require_firebase_uid()
            except PermissionError as exc:
                return render_home(error=str(exc)), 401
            from resume_profiles import load_resume_profile, profile_is_current

            cached_resume_profile = load_resume_profile(firebase_uid)
            if cached_resume_profile is None:
                return render_home(error="No saved resume profile was found. Upload and process one first."), 400
            if not profile_is_current(cached_resume_profile):
                return render_home(error="Your saved resume profile is out of date. Please process your resume again."), 409
            resume_text = ""

        if not resume_text and cached_resume_profile is None:
            return render_home(
                error="Upload a PDF/DOCX resume or paste resume text.",
                country=country,
                state=state,
                max_required_yoe=max_required_yoe,
                exclude_security_clearance=exclude_security_clearance,
                require_recent_posted=require_recent_posted,
            ), 400

        if len(resume_text) > MAX_RESUME_CHARS:
            return render_home(
                error=f"Resume text is too long. Please keep it under {MAX_RESUME_CHARS:,} characters.",
                resume_text=resume_text[:MAX_RESUME_CHARS],
                country=country,
                state=state,
                max_required_yoe=max_required_yoe,
                exclude_security_clearance=exclude_security_clearance,
                require_recent_posted=require_recent_posted,
            ), 413

        # Heavy imports happen only after the user submits the form.
        # This keeps the homepage fast.
        from ranking_service import run_resume_ranking

        payload: dict[str, Any] = run_resume_ranking(
            resume_text=resume_text,
            country=country,
            state=state,
            max_required_yoe=max_required_yoe,
            exclude_security_clearance=exclude_security_clearance,
            require_recent_posted=require_recent_posted,
            top_k_to_show=TOP_K_TO_SHOW,
            precomputed_resume_profile=cached_resume_profile,
        )

        return render_template(
            "results.html",
            payload=payload,
            country=country,
            state=state,
            max_required_yoe=max_required_yoe,
            enable_max_yoe_filter=ENABLE_MAX_YOE_FILTER,
            require_recent_posted=require_recent_posted,
            top_k_to_show=TOP_K_TO_SHOW,
            show_ranking_diagnostics=SHOW_RANKING_DIAGNOSTICS,
        )

    except HTTPException:
        raise
    except ResumeExtractionError as exc:
        return render_home(
            error=str(exc),
            resume_text=request.form.get("resume_text", "").strip(),
            country=request.form.get("country", "").strip() or None,
            state=request.form.get("state", "").strip() or None,
            max_required_yoe=parse_float_or_none(request.form.get("max_required_yoe")) if ENABLE_MAX_YOE_FILTER else None,
            exclude_security_clearance=request.form.get("exclude_security_clearance") == "on",
            require_recent_posted=request.form.get("require_recent_posted") == "on",
        ), 400
    except Exception as exc:
        log_safe_request_error("/rank", exc)

        return render_home(
            error="Something went wrong while ranking this resume. Please try again.",
        ), 500


@app.get("/resume-profile")
def resume_profile_status() -> Any:
    try:
        firebase_uid = require_firebase_uid()
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 401

    from resume_profiles import load_resume_profile, profile_is_current

    profile = load_resume_profile(firebase_uid)
    if profile is None:
        return jsonify({"exists": False})
    return jsonify(
        {
            "exists": True,
            "current": profile_is_current(profile),
            "updated_at": profile.get("updated_at"),
            "technologies": profile.get("technologies", []),
            "technology_categories": profile.get("technology_categories", []),
        }
    )


@app.post("/resume-profile")
@limiter.limit("3 per minute")
def create_resume_profile() -> Any:
    try:
        firebase_uid = require_firebase_uid()
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 401

    try:
        resume_text, *_ = get_rank_form_inputs()
        if not resume_text:
            return jsonify({"error": "Upload a PDF/DOCX resume or paste resume text."}), 400
        if len(resume_text) > MAX_RESUME_CHARS:
            return jsonify({"error": f"Resume text must be under {MAX_RESUME_CHARS:,} characters."}), 413

        from ranking_service import get_minilm_model
        from resume_profiles import build_resume_profile, save_resume_profile

        profile = build_resume_profile(resume_text, get_minilm_model())
        save_resume_profile(firebase_uid, profile)
        return jsonify(
            {
                "saved": True,
                "technologies": profile.get("technologies", []),
                "technology_categories": profile.get("technology_categories", []),
            }
        )
    except ResumeExtractionError as exc:
        return jsonify({"error": str(exc)}), 400
    except HTTPException:
        raise
    except Exception as exc:
        log_safe_request_error("/resume-profile", exc)
        return jsonify({"error": "Could not process and save this resume profile."}), 500


@app.delete("/resume-profile")
def remove_resume_profile() -> Any:
    try:
        firebase_uid = require_firebase_uid()
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 401

    from resume_profiles import delete_resume_profile

    return jsonify({"deleted": delete_resume_profile(firebase_uid)})


@app.delete("/account")
@limiter.limit("3 per minute")
def delete_account() -> Any:
    """Delete the verified user's cached profile and Firebase account."""
    try:
        firebase_uid = require_firebase_uid()
    except PermissionError as exc:
        return jsonify({"error": str(exc)}), 401

    decoded_token = getattr(g, "firebase_user", None) or {}
    auth_time = decoded_token.get("auth_time")
    try:
        auth_age_seconds = time.time() - float(auth_time)
    except (TypeError, ValueError):
        auth_age_seconds = ACCOUNT_DELETE_RECENT_AUTH_SECONDS + 1

    if auth_age_seconds < -60 or auth_age_seconds > ACCOUNT_DELETE_RECENT_AUTH_SECONDS:
        return jsonify(
            {
                "error": "For security, sign out and sign back in before deleting your account.",
                "code": "recent_login_required",
            }
        ), 403

    from resume_profiles import delete_resume_profile

    try:
        delete_resume_profile(firebase_uid)
    except Exception as exc:
        log_safe_request_error("/account/profile-delete", exc)
        return jsonify({"error": "Could not delete your saved data. Your account was not deleted."}), 500

    try:
        delete_firebase_user(firebase_uid)
    except Exception as exc:
        log_safe_request_error("/account/firebase-delete", exc)
        return jsonify(
            {
                "error": "Your saved profile was deleted, but the account could not be deleted. Please try again.",
            }
        ), 500

    return jsonify({"deleted": True})


@app.get("/about")
def about() -> str:
    """
    Show the about page.
    """
    return render_template("about.html")


@app.get("/privacy")
def privacy() -> str:
    """
    Show the privacy policy page.
    """
    return render_template("privacy.html")


@app.get("/terms")
def terms() -> str:
    """
    Show the terms of service page.
    """
    return render_template("terms.html")


@app.get("/robots.txt")
def robots_txt() -> Any:
    """
    Serve crawler directives from the root URL expected by crawlers.
    """
    return send_from_directory(app.static_folder, "robots.txt", mimetype="text/plain")


@app.get("/sitemap.xml")
def sitemap_xml() -> Any:
    """
    Serve the public page sitemap from the root URL expected by crawlers.
    """
    return send_from_directory(app.static_folder, "sitemap.xml", mimetype="application/xml")


@app.get("/health")
def health() -> tuple[str, int]:
    """
    Simple health check endpoint for Cloud Run.
    """
    return "ok", 200


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "8080"))
    app.run(host="0.0.0.0", port=port)
